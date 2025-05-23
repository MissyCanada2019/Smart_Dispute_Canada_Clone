import os
from docx import Document
from datetime import datetime

def generate_docx(case, user):
    """
    Generates a basic DOCX file with user and case details.
    Returns the full file path to the generated document.
    """
    filename = f"{user.full_name.replace(' ', '_')}_case_{case.id}.docx"
    output_dir = "/tmp/generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading('Legal Case Summary', level=1)
    doc.add_paragraph(f"Name: {user.full_name}")
    doc.add_paragraph(f"Email: {user.email}")
    doc.add_paragraph(f"Case Title: {case.title}")
    doc.add_paragraph(f"Description: {case.description or 'N/A'}")
    doc.add_paragraph(f"Date Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}")

    doc.save(filepath)
    return filepath
