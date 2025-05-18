import os
from docx import Document
from datetime import datetime

TEMPLATE_DIR = "templates/forms"

def generate_docx(case, user, form_type="repair_request"):
    """
    Generate a DOCX form filled with user and case info based on province.
    """
    province = user.province.lower() if hasattr(user, 'province') else "on"
    template_path = os.path.join(TEMPLATE_DIR, province, f"{form_type}_template.docx")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # Replace basic placeholders
    placeholders = {
        "[FULL_NAME]": user.full_name,
        "[EMAIL]": user.email,
        "[LEGAL_ISSUE]": case.legal_issue,
        "[DATE]": datetime.utcnow().strftime("%Y-%m-%d")
    }

    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Add conditional clauses based on legal issue
    issue_text = case.legal_issue.lower()
    if "mold" in issue_text:
        doc.add_paragraph("This case involves mold, which raises significant habitability concerns under the Residential Tenancies Act.")
    if "eviction" in issue_text:
        doc.add_paragraph("The tenant may be facing an unlawful eviction, requiring urgent legal remedy.")
    if "harassment" in issue_text:
        doc.add_paragraph("The tenant reports sustained harassment, affecting quiet enjoyment of the rental unit.")

    # Save file
    output_dir = "generated_forms"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f"{province}_{form_type}_case{case.id}.docx")
    doc.save(output_path)

    return output_path
