import os
from docx import Document
from datetime import datetime

TEMPLATE_DIR = "templates/forms"
OUTPUT_DIR = "generated_forms"

def generate_docx(case, user, form_type="repair_request"):
    province = user.province.lower()
    template_path = os.path.join(TEMPLATE_DIR, province, f"{form_type}_template.docx")

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    placeholders = {
        "[FULL_NAME]": user.full_name,
        "[EMAIL]": user.email,
        "[LEGAL_ISSUE]": case.legal_issue,
        "[DATE]": datetime.utcnow().strftime("%Y-%m-%d")
    }

    for para in doc.paragraphs:
        for ph, val in placeholders.items():
            if ph in para.text:
                para.text = para.text.replace(ph, val)

    issue_text = case.legal_issue.lower()
    if "mold" in issue_text:
        doc.add_paragraph("This case involves mold, which raises habitability concerns.")
    if "eviction" in issue_text:
        doc.add_paragraph("Potential unlawful eviction in breach of tenancy law.")
    if "harassment" in issue_text:
        doc.add_paragraph("Tenant reports harassment affecting quiet enjoyment.")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(OUTPUT_DIR, f"{province}_{form_type}_case{case.id}.docx")
    doc.save(docx_path)

    return docx_path

# Confirm utility setup
os.path.exists(TEMPLATE_DIR), os.path.exists(OUTPUT_DIR)
